"""
docx_okuyucu ve cikti_ureteci modülleri için testler.
"""

import os
import sys
import tempfile
import unittest

from docx import Document

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx_okuyucu import docx_oku
from cikti_ureteci import cikti_olustur


def _tablo_docx_olustur(dosya_yolu: str, satirlar: list[tuple[str, str]]) -> None:
    """Verilen etiket-değer çiftlerinden tablolu bir .docx dosyası oluşturur."""
    belge = Document()
    tablo = belge.add_table(rows=0, cols=2)
    for etiket, deger in satirlar:
        satir = tablo.add_row()
        satir.cells[0].text = etiket
        satir.cells[1].text = deger
    belge.save(dosya_yolu)


def _paragraf_docx_olustur(dosya_yolu: str, satirlar: list[str]) -> None:
    """Verilen paragraf listesinden bir .docx dosyası oluşturur."""
    belge = Document()
    for satir in satirlar:
        belge.add_paragraph(satir)
    belge.save(dosya_yolu)


class TestDocxOkuyucu(unittest.TestCase):
    def test_tablodan_okuma(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            dosya = tmp.name

        try:
            _tablo_docx_olustur(dosya, [
                ("Ders Adı", "Matematik"),
                ("Sınıf / Şube", "7 / A"),
                ("Konu", "Kesirler"),
                ("Süre", "40 dakika"),
            ])

            veri = docx_oku(dosya)

            self.assertEqual(veri.get("ders_adi"), "Matematik")
            self.assertEqual(veri.get("sinif_sube"), "7 / A")
            self.assertEqual(veri.get("konu"), "Kesirler")
            self.assertEqual(veri.get("sure"), "40 dakika")
        finally:
            os.unlink(dosya)

    def test_paragraflardan_okuma(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            dosya = tmp.name

        try:
            _paragraf_docx_olustur(dosya, [
                "Ders Adı: Fen Bilimleri",
                "Konu: Hücre Bölünmesi",
                "Süre: 40 dakika",
            ])

            veri = docx_oku(dosya)

            self.assertEqual(veri.get("ders_adi"), "Fen Bilimleri")
            self.assertEqual(veri.get("konu"), "Hücre Bölünmesi")
            self.assertEqual(veri.get("sure"), "40 dakika")
        finally:
            os.unlink(dosya)

    def test_dosya_bulunamadi(self):
        with self.assertRaises(FileNotFoundError):
            dosya = os.path.join(tempfile.gettempdir(), "var_olmayan_dosya_xyz.docx")
            docx_oku(dosya)

    def test_bos_icerik_hatasi(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            dosya = tmp.name

        try:
            belge = Document()
            belge.add_paragraph("Herhangi bir etiket içermeyen metin.")
            belge.save(dosya)

            with self.assertRaises(ValueError):
                docx_oku(dosya)
        finally:
            os.unlink(dosya)


class TestCiktiUreteci(unittest.TestCase):
    def test_cikti_olusturma(self):
        veri = {
            "ders_adi": "Türkçe",
            "sinif_sube": "5 / B",
            "konu": "Noktalama İşaretleri",
            "sure": "40 dakika",
            "kazanimlar": "Noktalama işaretlerini doğru kullanır.",
        }

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            cikti_dosya = tmp.name

        try:
            cikti_olustur(veri, cikti_dosya)

            self.assertTrue(os.path.exists(cikti_dosya))
            self.assertGreater(os.path.getsize(cikti_dosya), 0)

            belge = Document(cikti_dosya)
            tum_metin = " ".join(
                hucre.text
                for tablo in belge.tables
                for satir in tablo.rows
                for hucre in satir.cells
            )
            self.assertIn("Türkçe", tum_metin)
            self.assertIn("Noktalama İşaretleri", tum_metin)
        finally:
            os.unlink(cikti_dosya)

    def test_bos_veri(self):
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            cikti_dosya = tmp.name

        try:
            cikti_olustur({}, cikti_dosya)
            self.assertTrue(os.path.exists(cikti_dosya))
        finally:
            os.unlink(cikti_dosya)


class TestMain(unittest.TestCase):
    def test_uctan_uca(self):
        """Örnek bir docx dosyasını okuyup çıktı oluşturur."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            girdi = tmp.name
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            cikti = tmp.name

        try:
            _tablo_docx_olustur(girdi, [
                ("Ders Adı", "Tarih"),
                ("Sınıf / Şube", "10 / C"),
                ("Konu", "Osmanlı Devleti'nin Kuruluşu"),
                ("Süre", "40 dakika"),
                ("Kazanımlar", "Osmanlı Devleti'nin kuruluş sürecini açıklar."),
                ("Materyaller", "Harita, ders kitabı"),
                ("Değerlendirme", "Sözlü sorular"),
            ])

            veri = docx_oku(girdi)
            cikti_olustur(veri, cikti)

            self.assertTrue(os.path.exists(cikti))
            belge = Document(cikti)
            tum_metin = " ".join(
                hucre.text
                for tablo in belge.tables
                for satir in tablo.rows
                for hucre in satir.cells
            )
            self.assertIn("Tarih", tum_metin)
            self.assertIn("Osmanlı Devleti'nin Kuruluşu", tum_metin)
        finally:
            os.unlink(girdi)
            os.unlink(cikti)


if __name__ == "__main__":
    unittest.main()
