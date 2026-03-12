"""
Çıktı Üreteci Modülü

Bu modül, ders planı verilerini alır ve biçimlendirilmiş bir .docx
dosyası olarak kaydeder.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


ALAN_ETIKETLERI = {
    "ders_adi": "Ders Adı",
    "sinif": "Sınıf",
    "sube": "Şube",
    "sinif_sube": "Sınıf / Şube",
    "konu": "Konu",
    "kazanimlar": "Kazanımlar",
    "sure": "Süre",
    "yontem_teknik": "Yöntem ve Teknikler",
    "materyaller": "Materyaller",
    "degerlendirme": "Değerlendirme",
}

ALAN_SIRASI = [
    "ders_adi",
    "sinif_sube",
    "sinif",
    "sube",
    "konu",
    "sure",
    "kazanimlar",
    "yontem_teknik",
    "materyaller",
    "degerlendirme",
]


def _baslik_ekle(belge: Document, metin: str) -> None:
    """Belgeye merkezi, büyük ve kalın başlık paragrafı ekler."""
    paragraf = belge.add_paragraph()
    paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraf.add_run(metin)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)


def _tablo_satiri_ekle(tablo, etiket: str, deger: str) -> None:
    """Tabloya iki sütunlu bir satır ekler: etiket (kalın) | değer."""
    satir = tablo.add_row()
    etiket_hucre = satir.cells[0]
    deger_hucre = satir.cells[1]

    etiket_run = etiket_hucre.paragraphs[0].add_run(etiket)
    etiket_run.bold = True
    etiket_run.font.size = Pt(11)

    deger_run = deger_hucre.paragraphs[0].add_run(deger)
    deger_run.font.size = Pt(11)


def cikti_olustur(veri: dict, cikti_yolu: str) -> None:
    """
    Ders planı verilerinden biçimlendirilmiş bir .docx dosyası oluşturur.

    Parametreler
    ----------
    veri : dict
        ``docx_okuyucu.docx_oku`` tarafından döndürülen ders planı verisi.
    cikti_yolu : str
        Oluşturulacak .docx dosyasının yolu.
    """
    belge = Document()

    _baslik_ekle(belge, "DERS PLANI")
    belge.add_paragraph()

    tablo = belge.add_table(rows=0, cols=2)
    tablo.style = "Table Grid"

    tablo.columns[0].width = Pt(120)

    for alan in ALAN_SIRASI:
        if alan in veri:
            etiket = ALAN_ETIKETLERI.get(alan, alan)
            _tablo_satiri_ekle(tablo, etiket, veri[alan])

    for alan, deger in veri.items():
        if alan not in ALAN_SIRASI:
            etiket = ALAN_ETIKETLERI.get(alan, alan)
            _tablo_satiri_ekle(tablo, etiket, deger)

    belge.save(cikti_yolu)
