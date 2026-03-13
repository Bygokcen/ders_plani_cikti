"""
ogretim_elemanlari.py
----------------------
Program_Kılavuzu_BILMUH2025-26_19022026.docx dosyasından
tüm öğretim elemanlarını fotoğraflarıyla birlikte çıkarır.

Çıktı:
  ogretim_elemanlari.json          -- tüm elemanların verisi
  ogretim_elemanlari_foto/<ad>.ext -- fotoğraf dosyaları
"""

import argparse
import docx
import json
import os
import re
from pathlib import Path

from lxml import etree

ROOT_DIR = Path(__file__).resolve().parents[1]
DEFAULT_JSON_OUT = Path(__file__).resolve().with_name('ogretim_elemanlari.json')
DEFAULT_IMG_BASE = Path(__file__).resolve().with_name('ogretim_elemanlari_foto')

WNS  = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A_NS = 'http://schemas.openxmlformats.org/drawingml/2006/main'
EMB  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
DRAW = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'


def slugify(name):
    """Ad-soyad'dan dosya adı üret."""
    name = name.strip()
    # Unvanları at (Prof. Dr., Doç. Dr., Dr. Öğr. Üyesi, Arş. Gör. Dr. vb.)
    name = re.sub(
        r'^(Prof\.|Doç\.|Dr\.|Öğr\.|Arş\.|Gör\.|Üyesi|Yrd\.)\s*', '', name
    ).strip()
    name = re.sub(r'\s+', '_', name)
    # Türkçe karakter dönüşümü
    tr_map = str.maketrans('çğıöşüÇĞİÖŞÜ', 'cgiosucgiosu')
    name = name.translate(tr_map)
    name = re.sub(r'[^A-Za-z0-9_-]', '', name)
    return name or 'ogretim_elemani'


def parse_cell_text(cell):
    """
    Hücre metnini satır satır okuyup alanlara ayırır.
    Yapı:
      Satır 1 : Ad Soyad (+ unvan baştaysa)
      Satır N : email@gop.edu.tr
      Satır N : İç Hat: XXXX
      Satır N : Çalışma Alanları: ...
    """
    paragraphs = [p.text.strip() for p in cell.paragraphs if p.text.strip()]

    ad_soyad       = ''
    eposta         = ''
    ic_hat         = ''
    calisma_alanlari = ''
    extra_lines    = []

    for line in paragraphs:
        if '@gop.edu.tr' in line or '@' in line:
            eposta = line.strip()
        elif line.startswith('İç Hat'):
            ic_hat = re.sub(r'İç Hat\s*[:\-]?\s*', '', line).strip()
        elif 'Çalışma Alanları' in line or 'Çalismus' in line:
            calisma_alanlari = re.sub(r'Çalışma Alanları\s*[:\-]?\s*', '', line).strip()
        elif not ad_soyad:
            ad_soyad = line
        else:
            extra_lines.append(line)

    # Bazen ad ve e-posta aynı satırda olabilir
    if not ad_soyad and eposta:
        ad_soyad = eposta

    return {
        'ad_soyad': ad_soyad,
        'eposta': eposta,
        'ic_hat': ic_hat,
        'calisma_alanlari': calisma_alanlari,
    }


def extract_photo(cell, doc_part, slug, img_base):
    """Hücredeki ilk gömülü resmi slug.ext adıyla diske yazar."""
    tc = cell._tc
    drawings = tc.findall(f'.//{DRAW}inline') + tc.findall(f'.//{DRAW}anchor')
    for drawing in drawings:
        blip = drawing.find(f'.//{{{A_NS}}}blip')
        if blip is None:
            continue
        r_embed = blip.get(f'{{{EMB}}}embed')
        if not r_embed:
            continue
        try:
            img_part  = doc_part.rels[r_embed].target_part
            img_bytes = img_part.blob
            ext       = img_part.partname.split('.')[-1].lower()
            os.makedirs(img_base, exist_ok=True)
            fname = f'{slug}.{ext}'
            fpath = os.path.join(img_base, fname)
            with open(fpath, 'wb') as f:
                f.write(img_bytes)
            # Proje köküne göreli yol döndür
            return os.path.relpath(fpath, ROOT_DIR)
        except Exception as e:
            return f'[hata: {e}]'
    return None


def find_default_docx() -> Path | None:
    candidates = sorted(ROOT_DIR.glob('*.docx'))
    if not candidates:
        return None
    for candidate in candidates:
        if 'program_kılavuzu' in candidate.name.lower() or 'program_kilavuzu' in candidate.name.lower():
            return candidate
    return candidates[0]


def raw_row_cells(row):
    return list(getattr(row._tr, 'tc_lst', []))


def raw_row_text(row) -> str:
    parts = []
    for tc in raw_row_cells(row):
        texts = [node.text or '' for node in tc.iter(f'{WNS}t')]
        joined = ' '.join(' '.join(texts).split())
        if joined:
            parts.append(joined)
    return ' '.join(parts)


def row_has_photo(row) -> bool:
    for tc in raw_row_cells(row):
        drawings = tc.findall(f'.//{DRAW}inline') + tc.findall(f'.//{DRAW}anchor')
        if drawings:
            return True
    return False


def faculty_table_score(tbl) -> int:
    email_hits = 0
    office_hits = 0
    research_hits = 0
    photo_hits = 0

    for row in tbl.rows:
        row_text = raw_row_text(row)
        if '@' in row_text:
            email_hits += 1
        if 'iç hat' in row_text.lower() or 'ic hat' in row_text.lower():
            office_hits += 1
        if 'çalışma alanları' in row_text.lower() or 'calisma alanlari' in row_text.lower() or 'çalismus' in row_text.lower():
            research_hits += 1
        if row_has_photo(row):
            photo_hits += 1

    if email_hits < 2:
        return -1

    return (research_hits * 10) + (photo_hits * 4) + (office_hits * 2) + email_hits + min(len(tbl.rows), 20)


def pick_faculty_table(doc, explicit_index: int | None):
    if explicit_index is not None:
        return doc.tables[explicit_index], explicit_index

    best_index = None
    best_score = -1
    for index, table in enumerate(doc.tables):
        score = faculty_table_score(table)
        if score > best_score:
            best_score = score
            best_index = index

    if best_index is None or best_score < 0:
        raise RuntimeError('Ogretim elemanlari tablosu otomatik bulunamadi. --table-index ile belirtin.')

    return doc.tables[best_index], best_index


def main():
    default_docx = find_default_docx()
    parser = argparse.ArgumentParser(description='DOCX\'ten ogretim elemanlarini ve fotograflarini cikarir')
    parser.add_argument('--docx', default=str(default_docx) if default_docx else None, help='Kaynak DOCX yolu')
    parser.add_argument('--out', default=str(DEFAULT_JSON_OUT), help='Cikti JSON yolu')
    parser.add_argument('--img-base', default=str(DEFAULT_IMG_BASE), help='Fotograflarin yazilacagi klasor')
    parser.add_argument('--table-index', type=int, default=None, help='Ogretim elemanlari tablosunun 0-based indexi')
    args = parser.parse_args()

    if not args.docx:
        raise FileNotFoundError('Kaynak DOCX bulunamadi. --docx ile dosya yolu verin.')

    docx_path = Path(args.docx)
    json_out = Path(args.out)
    img_base = Path(args.img_base)

    doc = docx.Document(str(docx_path))
    tbl, table_index = pick_faculty_table(doc, args.table_index)

    results  = {}
    counters = {}  # aynı slug varsa numara ekle

    for row in tbl.rows:
        if not row.cells:
            continue
        cell = row.cells[0]
        info = parse_cell_text(cell)

        if not info['ad_soyad']:
            continue

        slug = slugify(info['ad_soyad'])
        # Çakışma varsa numara ekle
        if slug in counters:
            counters[slug] += 1
            slug = f'{slug}_{counters[slug]}'
        else:
            counters[slug] = 1

        foto = extract_photo(cell, doc.part, slug, str(img_base))
        info['foto'] = foto

        results[slug] = info
        print(f"  {info['ad_soyad']:50s} -> foto: {foto}")

    json_out.parent.mkdir(parents=True, exist_ok=True)
    with open(json_out, 'w', encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    print(f"\nToplam: {len(results)} öğretim elemanı")
    print(f"Tablo : {table_index}")
    print(f"JSON  : {json_out}")
    foto_count = sum(1 for v in results.values() if v['foto'] and not v['foto'].startswith('['))
    print(f"Fotoğraf: {foto_count} dosya ({img_base})")


if __name__ == '__main__':
    main()
