"""
Örnek ders planı .docx dosyası oluşturur.
Projenin doğru çalıştığını doğrulamak için kullanılabilir.

Kullanım:
    python ornek_olustur.py
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ORNEK_VERI = {
    "Ders Adı": "Matematik",
    "Sınıf / Şube": "7 / A",
    "Konu": "Kesirlerle Çarpma İşlemi",
    "Kazanımlar": (
        "• Kesirlerle çarpma işlemi yapabilir.\n"
        "• Sonucu sadeleştirerek sade kesir olarak ifade edebilir."
    ),
    "Süre": "40 dakika",
    "Yöntem ve Teknikler": "Soru-cevap, grup çalışması, problem çözme",
    "Materyaller": "Ders kitabı, tahta, renkli kartlar",
    "Değerlendirme": "Sözlü soru-cevap, çalışma kâğıdı",
}


def ornek_olustur(dosya_yolu: str = "ornek_ders_plani.docx") -> None:
    belge = Document()

    baslik = belge.add_paragraph("DERS PLANI")
    baslik.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = baslik.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    belge.add_paragraph()

    tablo = belge.add_table(rows=0, cols=2)
    tablo.style = "Table Grid"

    for etiket, deger in ORNEK_VERI.items():
        satir = tablo.add_row()
        satir.cells[0].text = etiket
        satir.cells[1].text = deger

    belge.save(dosya_yolu)
    print(f"Örnek dosya oluşturuldu: {dosya_yolu}")


if __name__ == "__main__":
    ornek_olustur()
