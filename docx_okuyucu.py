"""
Docx Okuyucu Modülü

Bu modül, ders planı içeren .docx dosyalarını okur ve yapılandırılmış
veri olarak döndürür.

Ders planı formatı (tablo veya etiketli paragraflar):
  - Ders Adı
  - Sınıf / Şube
  - Konu
  - Kazanımlar
  - Süre
  - Yöntem ve Teknikler
  - Materyaller
  - Değerlendirme
"""

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

ALAN_ETIKETLERI = {
    "ders adı": "ders_adi",
    "sınıf": "sinif",
    "şube": "sube",
    "sınıf / şube": "sinif_sube",
    "konu": "konu",
    "kazanımlar": "kazanimlar",
    "süre": "sure",
    "yöntem ve teknikler": "yontem_teknik",
    "materyaller": "materyaller",
    "değerlendirme": "degerlendirme",
}


def _normalize(metin: str) -> str:
    """Metni küçük harfe çevirir ve baştaki/sondaki boşlukları temizler."""
    return metin.strip().lower()


def _tablodan_oku(tablo) -> dict:
    """Bir docx tablosundan etiket → değer çiftlerini çıkarır."""
    veri = {}
    for satir in tablo.rows:
        hucreler = satir.cells
        if len(hucreler) < 2:
            continue
        etiket = _normalize(hucreler[0].text)
        deger = hucreler[1].text.strip()
        anahtar = ALAN_ETIKETLERI.get(etiket)
        if anahtar:
            veri[anahtar] = deger
    return veri


def _paragraflardan_oku(paragraflar) -> dict:
    """
    Etiketli paragraflardan ders planı alanlarını çıkarır.
    Her alan "Etiket: Değer" ya da "Etiket:" biçiminde başlayan
    bir paragrafla başlar; sonraki satırlar o alanın devamıdır.
    """
    veri = {}
    mevcut_anahtar = None
    mevcut_satirlar = []

    def _kaydet():
        if mevcut_anahtar:
            veri[mevcut_anahtar] = "\n".join(mevcut_satirlar).strip()

    for paragraf in paragraflar:
        metin = paragraf.text.strip()
        if not metin:
            continue

        eslesme = None
        for etiket, anahtar in ALAN_ETIKETLERI.items():
            if metin.lower().startswith(etiket + ":"):
                eslesme = (anahtar, metin[len(etiket) + 1:].strip())
                break

        if eslesme:
            _kaydet()
            mevcut_anahtar, ilk_deger = eslesme
            mevcut_satirlar = [ilk_deger] if ilk_deger else []
        elif mevcut_anahtar:
            mevcut_satirlar.append(metin)

    _kaydet()
    return veri


def docx_oku(dosya_yolu: str) -> dict:
    """
    Belirtilen .docx dosyasını okur ve ders planı verilerini döndürür.

    Parametreler
    ----------
    dosya_yolu : str
        Okunacak .docx dosyasının yolu.

    Döndürür
    --------
    dict
        Ders planı alanlarını içeren sözlük.

    Hata
    ----
    FileNotFoundError
        Dosya bulunamazsa.
    ValueError
        Dosyadan hiçbir ders planı alanı çıkarılamazsa.
    """
    try:
        belge = Document(dosya_yolu)
    except PackageNotFoundError:
        raise FileNotFoundError(f"'{dosya_yolu}' dosyası bulunamadı.")

    veri = {}

    for tablo in belge.tables:
        veri.update(_tablodan_oku(tablo))

    if not veri:
        veri.update(_paragraflardan_oku(belge.paragraphs))

    if not veri:
        raise ValueError(
            f"'{dosya_yolu}' dosyasından hiçbir ders planı alanı çıkarılamadı. "
            "Lütfen dosyanın beklenen formatta olduğunu doğrulayın."
        )

    return veri
